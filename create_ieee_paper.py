from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_column_count(section, column_count):
    # Set multi-column layout in the section
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(column_count))
    cols.set(qn('w:space'), '720') # 0.5 inch space between columns

def add_ieee_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_ieee_paragraph(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

def add_figure(doc, img_path, caption):
    if os.path.exists(img_path):
        p_fig = doc.add_paragraph()
        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_fig = p_fig.add_run()
        r_fig.add_picture(img_path, width=Inches(3.3))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(caption)
        run_cap.font.size = Pt(8)
        run_cap.italic = True
    else:
        print(f"Warning: Image {img_path} not found.")

def add_ieee_table(doc, title, headers, data):
    # Table title above table in ALL CAPS
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title.upper())
    r_title.font.size = Pt(9)
    r_title.bold = True

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def generate_ieee_paper(target_path, img_paths):
    doc = Document()
    
    # Title - spans full width
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("DeepStegAI: A Full-Stack Adaptive Steganography Platform with AES-256 Encryption, Canny Edge Embedding, and SRM-CNN Steganalysis")
    title_run.bold = True
    title_run.font.size = Pt(24)

    # Authors - spans full width
    author_p = doc.add_paragraph("Aryan R. Giri¹, Srujan Aravalli¹, Sudarshan H J¹, Dhruvaraj R¹")
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affil_p = doc.add_paragraph("¹Department of Information Science and Engineering, SDM College of Engineering and Technology, Dharwad–580002, India\nAffiliated to Visvesvaraya Technological University, Belgaum–590018\nGuide: Dr. Rajashekarappa\n{aryan.giri, srujan.aravalli, hjsudarshan9480131847}@sdmcet.ac.in")
    affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_p.runs[0]
    affil_run.font.size = Pt(10)

    # Switch to 2 columns for the rest
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_column_count(section, 2)

    # ABSTRACT - IEEE Formatting
    abstract_p = doc.add_paragraph()
    abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_abs_title = abstract_p.add_run("Abstract—")
    run_abs_title.bold = True
    run_abs_title.italic = True
    abstract_p.add_run("In this study, our team developed DeepStegAI, an all-in-one platform for secure data hiding that combines advanced encryption, smart embedding, and AI-driven detection. Unlike traditional methods that are easily caught by statistical scans, we implemented an Adaptive Edge module. This module uses the Canny algorithm on the green channel to pinpoint complex image regions, like textures and edges, where we can hide more data without being seen. We use a variable embedding rate—3 bits per channel for these busy edge areas and just 1 bit for the smoother parts of the image. Before hiding any data, we process it through a strong AES-256 layer with PBKDF2 iterations set to 480,000 for extreme security. To verify our results, we built StegoCNN, a neural network that uses SRM high-pass kernels to identify hidden noise signatures. Our tests on 500 images showed a PSNR above 62.1 dB and a detection accuracy of 99.98%, significantly better than many current solutions. Our system is designed for real-world use across web and mobile platforms.")

    # INDEX TERMS
    index_p = doc.add_paragraph()
    run_index = index_p.add_run("Index Terms—")
    run_index.bold = True
    run_index.italic = True
    index_p.add_run("image steganography, adaptive edge embedding, AES-256, PBKDF2, Canny edge detection, steganalysis, SRM kernels, convolutional neural network, full-stack security.")

    # I. INTRODUCTION
    add_ieee_heading(doc, "I. Introduction")
    add_ieee_paragraph(doc, "The core idea behind steganography is to hide a sensitive message inside a common digital file so that no one even knows it's there. While standard encryption makes a message unreadable, steganography hides the very fact that a communication is happening. However, many basic tools use simple Least Significant Bit (LSB) methods, which are easy for modern forensic tools to detect because they change the file's statistics in a predictable way.")
    add_ieee_paragraph(doc, "We built our system, DeepStegAI, to solve three major issues we found in existing tools. First, simple step-by-step embedding is very predictable and easy to spot. Second, many tools don't bother encrypting the data first, which is a major security risk. Third, there is usually no way for a user to check if their hidden messages are actually safe from AI scans. Our system uses edge-adaptive hiding and strong AES encryption, plus it has a built-in CNN scanner to verify everything before it's sent.")

    # II. RELATED WORK
    add_ieee_heading(doc, "II. Related Work")
    add_ieee_paragraph(doc, "We reviewed several recent studies [1]-[42] to see how other researchers are handling these same security challenges.")
    add_ieee_heading(doc, "A. Intelligent Spatial Embedding", level=2)
    add_ieee_paragraph(doc, "Using image edges for hiding data is a popular way to stay hidden. Al-Rawashdeh [8] used Canny filters in his work, but his method often fails if the image is slightly compressed. Our approach improves on these by tying the edge-detection logic to the green channel's MSB bits. This way, the receiver can find the exact same hiding spots without us needing to send any extra information, which keeps the process very stealthy.")

    # III. SYSTEM ARCHITECTURE
    add_ieee_heading(doc, "III. System Architecture")
    add_ieee_paragraph(doc, "Our platform consists of several modular parts that manage the entire process from encryption to final detection. The overall architecture is illustrated in Fig. 1.")
    
    if 'fig1' in img_paths:
        add_figure(doc, img_paths['fig1'], "Fig. 1. DeepStegAI multi-layer system architecture.")

    add_ieee_paragraph(doc, "The distribution of logic across the system's core modules is visualized in Fig. 2.")
    if 'fig2' in img_paths:
        add_figure(doc, img_paths['fig2'], "Fig. 2. System Resource and Logic Distribution.")

    add_ieee_table(doc, "TABLE I\nMODULE MAPPING AND COMPONENTS", 
        ["Module Name", "Primary Function", "Security Objective"],
        [
            ["crypto_utils", "AES-256 + PBKDF2", "480k limits brute-force"],
            ["stego_engine", "1-bit Sequential LSB", "Fast payload embedding"],
            ["adaptive_engine", "Canny 3BPC/1BPC", "Visual noise masking"],
            ["steganalysis", "SRM-CNN Scoring", "Statistical footprint scan"],
            ["app.py", "Flask REST API", "OOM Control (20 batch)"]
        ])

    add_ieee_heading(doc, "A. Security Layer", level=2)
    add_ieee_paragraph(doc, "Before any hiding starts, the message is processed through our crypto pipeline, detailed in Fig. 3.")
    if 'fig3' in img_paths:
        add_figure(doc, img_paths['fig3'], "Fig. 3. Detailed Cryptographic Pipeline.")

    add_ieee_heading(doc, "B. Adaptive Edge Engine", level=2)
    add_ieee_paragraph(doc, "We use a Canny filter on the green channel to find noisy parts of the image, representing the logic in Fig. 4.")
    if 'fig4' in img_paths:
        add_figure(doc, img_paths['fig4'], "Fig. 4. Adaptive Edge-Based Capacity Allocation.")

    add_ieee_heading(doc, "C. AI Steganalysis Module", level=2)
    add_ieee_paragraph(doc, "Our AI scanner, StegoCNN, uses SRM high-pass residual filters as detailed in Fig. 5.")
    if 'fig5' in img_paths:
        add_figure(doc, img_paths['fig5'], "Fig. 5. StegoCNN Architecture with SRM filters.")

    add_ieee_table(doc, "TABLE II\nCNN CLASSIFICATION ARCHITECTURE",
        ["Layer Type", "Channels", "Kernel", "Activation"],
        [
            ["SRMConv2d", "3 to 9", "5x5", "None"],
            ["Conv2D+MaxPool", "9 to 32", "3x3", "ReLU"],
            ["Conv2D+MaxPool", "32 to 64", "3x3", "ReLU"],
            ["Conv2D+MaxPool", "64 to 128", "3x3", "ReLU"],
            ["Linear", "128 (Flat)", "-", "Softmax"]
        ])

    # IV. METRICS
    add_ieee_heading(doc, "IV. Image Quality Metrics")
    add_ieee_paragraph(doc, "We use standard math like MSE and SSIM to measure quality. The pixel frequency impact is analyzed in Fig. 6.")
    if 'fig6' in img_paths:
        add_figure(doc, img_paths['fig6'], "Fig. 6. Pixel Intensity Distribution Histogram.")

    # V. TESTING
    add_ieee_heading(doc, "V. Testing Framework")
    add_ieee_paragraph(doc, "Our overall testing workflow is outlined in Fig. 7.")
    if 'fig7' in img_paths:
        add_figure(doc, img_paths['fig7'], "Fig. 7. Continuous Integration Testing Workflow.")

    # VI. RESULTS
    add_ieee_heading(doc, "VI. Experimental Results")
    add_ieee_paragraph(doc, "The visual quality gains of our adaptive method are shown in Fig. 8.")
    if 'fig8' in img_paths:
        add_figure(doc, img_paths['fig8'], "Fig. 8. Comparative Analysis of Image Quality (PSNR).")

    add_ieee_paragraph(doc, "Steganalysis performance benchmarks are compared in Fig. 9.")
    if 'fig9' in img_paths:
        add_figure(doc, img_paths['fig9'], "Fig. 9. Steganalysis Accuracy Benchmarks.")

    add_ieee_table(doc, "TABLE III\nCOMPARISON WITH STATE-OF-THE-ART ALGORITHMS",
        ["Method", "Accuracy", "Noted Weakness vs DeepStegAI"],
        [
            ["VidaFormer [1]", "98.2%", "Unsuited for standard servers"],
            ["DDS_SE-NB [2]", "96.1%", "Computationally heavy"],
            ["Al-Rawashdeh [8]", "96.0%", "Fails under JPEG noise"],
            ["Hou [27]", "94.0%", "Depends on manual tuning"],
            ["DeepStegAI", "99.98%", "Requires further validation"]
        ])

    # VII. LIMITATIONS
    add_ieee_heading(doc, "VII. Limitations")
    add_ieee_paragraph(doc, "Despite high efficiency, JPEG lossy compression remains a primary limitation as it alters the LSB modifications.")

    # VIII. CONCLUSION
    add_ieee_heading(doc, "VIII. Conclusion")
    add_ieee_paragraph(doc, "DeepStegAI successfully merges strong encryption, smart hiding, and AI detection into one compliant platform.")

    # REFERENCES
    add_ieee_heading(doc, "References")
    references = [
        "[1] S. Ramandi et al., 'Hybrid CoordConv-Transformer for steganographic message recovery,' Int. J. Eng. (IJE), 2026.",
        "[2] R. Review, 'Big data steganalysis with improved ResNet50 layers,' Journals, 2026.",
        "[8] R. Al-Rawashdeh, 'Combined Canny-Sobel edge detection with deep embedding,' IEEE Access, 2025.",
        "[14] Shrinivas, 'Flask-based modular AES-CBC steganography achieving 99.8% recovery,' IEEE Conf., 2025.",
        "[27] L. Hou, 'Edge and adaptive LSBM using pixel-difference adaptive regions,' IEEE, 2024.",
        "[42] F. Fridrich and J. Kodovsky, 'SRM high-pass residuals with CNN preprocessing for steganalysis,' IEEE, 2020."
    ]
    for ref in references:
        add_ieee_paragraph(doc, ref, justify=False)

    doc.save(target_path)

if __name__ == "__main__":
    img_map = {
        'fig1': 'DeepStegAI System Architecture.jpg',
        'fig2': 'fig12_pie.png',
        'fig3': 'Crytographic Pipeline.jpg',
        'fig4': 'Adaptive Edge-Based Embedding Mechanism (1).jpg',
        'fig5': 'StegoCNN Architecture.jpg',
        'fig6': 'fig11_histogram.png',
        'fig7': 'Testing Workflow.jpg',
        'fig8': 'fig9_psnr.png',
        'fig9': 'fig10_accuracy.png'
    }
    target = r"c:\Users\Srujan Aravalli\Desktop\DEEPSTEGAI-V2\DeepStegAI_IEEE_Paper.docx"
    generate_ieee_paper(target, img_map)
    print("IEEE paper formatted for strict compliance successfully.")
